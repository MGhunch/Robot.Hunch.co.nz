"""
ROBOT — READERS
===============
One door, three readers. Anything a client drops turns into words before
it joins the dump, so the dump stays what it has always been: a string.

That's the whole design decision. The obvious way to accept a screenshot
is to make the dump multimodal — but the dump is read by the FEEDER and by
EXTRACT, and going multimodal changes the contract for both. Reading the
image into words at the door means nothing downstream changes at all. It's
the same move EXTRACT makes: a favour at the threshold.

  .txt .md .csv        the browser already reads these; here for completeness
  .docx                python-docx, which the parcel already uses to write
  .pdf                 pdfplumber. A scanned PDF has no text layer and
                       falls through to the image reader.
  .png .jpg .webp .gif transcribed by a model — see prompts/reader.md

Nothing here calls a model. The image case returns needs_model=True and
copy_stage does that part, so this file stays cheap to import and test.
"""

import io
import os
import re

MAX_BYTES = 10 * 1024 * 1024        # a 10MB drop is already a lot of dump
MAX_CHARS = 12000                   # matches the browser's own slice

TEXT_EXT = (".txt", ".md", ".csv", ".markdown", ".text", ".log", ".json")
DOCX_EXT = (".docx",)
PDF_EXT = (".pdf",)
IMAGE_EXT = (".png", ".jpg", ".jpeg", ".webp", ".gif")

# What the door says it takes, for the file input and the failure line.
ACCEPT = ",".join(TEXT_EXT + DOCX_EXT + PDF_EXT + IMAGE_EXT)
SPOKEN = "Word, PDF, a screenshot, or plain text"

IMAGE_MIME = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
              ".webp": "image/webp", ".gif": "image/gif"}


def kind(name):
    ext = os.path.splitext(name or "")[1].lower()
    if ext in TEXT_EXT:
        return "text"
    if ext in DOCX_EXT:
        return "docx"
    if ext in PDF_EXT:
        return "pdf"
    if ext in IMAGE_EXT:
        return "image"
    return None


def _tidy(text):
    """Extractors leave a lot of air. Collapse it — the dump is read by a
    model that pays by the token, and blank lines aren't information."""
    text = text.replace("\r\n", "\n").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()[:MAX_CHARS]


def _read_docx(data):
    from docx import Document
    doc = Document(io.BytesIO(data))
    out = [p.text for p in doc.paragraphs]
    # Briefs live in tables more often than anyone admits.
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append(" · ".join(c for c in cells if c))
    return _tidy("\n".join(out))


def _read_pdf(data):
    import pdfplumber
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages[:40]:
            pages.append(page.extract_text() or "")
    return _tidy("\n\n".join(pages))


def read(name, data):
    """(text, needs_model, error). needs_model means 'this is an image, or
    a PDF with no text in it — hand it to the model reader'."""
    if not data:
        return "", False, "empty file"
    if len(data) > MAX_BYTES:
        return "", False, "too big — 10MB is the limit"
    k = kind(name)
    if k is None:
        return "", False, f"not a format I read. {SPOKEN}."
    try:
        if k == "text":
            return _tidy(data.decode("utf-8", "replace")), False, None
        if k == "docx":
            text = _read_docx(data)
            return (text, False, None) if text else ("", False, "nothing in it I could read")
        if k == "pdf":
            text = _read_pdf(data)
            # No text layer: it's a scan, which is an image wearing a PDF.
            return (text, False, None) if len(text) > 40 else ("", True, None)
        if k == "image":
            return "", True, None
    except Exception as e:
        return "", False, f"couldn't open it ({type(e).__name__})"
    return "", False, "not a format I read"


def image_block(name, data):
    """The content block for a model call, when it's the model's turn."""
    import base64
    ext = os.path.splitext(name or "")[1].lower()
    return {"type": "image", "source": {
        "type": "base64", "media_type": IMAGE_MIME.get(ext, "image/png"),
        "data": base64.standard_b64encode(data).decode("ascii")}}
