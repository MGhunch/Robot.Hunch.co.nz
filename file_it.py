"""
FILE IT
=======
The takeaway counter. Four fillings — PDF, DOC, PICS, CODE — and one
button. The page is a dumb renderer: this module says which fillings are
on the menu for this container and this run, builds the ones that are
ticked, and hands back links. No model. Nothing per-container in here —
the menu is read off spec.md (Outputs, IMAGES) and the run's uploads.

  POST /api/fillings   {container, run}             -> the four tiles
  POST /api/wrap       {container, run, form, chosen, copy, take, tweaks}
                                                     -> files built, links
  GET  /api/wrap/<run>/<name>                        -> the file

Files land under ROBOT_WRAPS/<run>/ — temp by nature, rebuilt on every
press of WRAP IT. Not the volume; the volume is for the log and the pics.

The copy doc is the Hunch copy deck, rendered from code rather than a
template: Hunch furniture — the deck's grid, type and red bars — with
HAI2 top left and the client's mark top right, COPY DOC in Bebas, a row
per module in spec order. HAI2 is static/hai2.png when it's there, words
until it is. The PDF is stubbed grey until its own project
lands (hit list 5, second half).
"""

from flask import Blueprint, jsonify, request, send_from_directory, session
from datetime import date
import io
import os
import re
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
HAI2 = next((os.path.join(STATIC_DIR, n) for n in ("hai2.png", "hai2.jpg") if os.path.isfile(os.path.join(STATIC_DIR, n))), None)

from auth import require_auth
import containers as CT
from engine import build_facts, render_copy, render_terms, TermsError

file_bp = Blueprint("file_it", __name__)

WRAPS_DIR = os.environ.get("ROBOT_WRAPS", os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "wraps"))
IMAGES_DIR = os.environ.get("ROBOT_IMAGES", os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "images"))

_SAFE = re.compile(r"[^a-z0-9-]")
_clean = lambda s, n=40: _SAFE.sub("", (s or "").lower().replace(" ", "-"))[:n]


# ---------------------------------------------------------------------------
# THE MENU — which fillings, and the line on a grey tile.
# Two reasons a tile goes grey: the container can't, or this run didn't.
# The lines are the robot's, in character, no sorry.
# ---------------------------------------------------------------------------

FILLINGS = [
    {"id": "pdf",  "title": "PDF",  "line": "Tidied up for sign off."},
    {"id": "doc",  "title": "DOC",  "line": "Every word, on the record."},
    {"id": "pics", "title": "PICS", "line": "All zipped up in a bag."},
    {"id": "code", "title": "CODE", "line": "Ready for the postman."},
]

def _emits(c, *words):
    """spec.md Outputs, 'files emitted' column — does any row name this?"""
    text = " ".join(o.get("files", "") for o in c["spec"]["outputs"]).lower()
    return any(w in text for w in words)

def _pics_in(run):
    folder = os.path.join(IMAGES_DIR, _clean(run))
    return sorted(f for f in os.listdir(folder) if re.fullmatch(r"\d+-[a-z0-9-]+\.(jpg|png)", f)) \
        if run and os.path.isdir(folder) else []

def menu(c, run):
    tiles = []
    for f in FILLINGS:
        t = dict(f, on=True)
        if f["id"] == "pdf":
            # its own project. Grey means not yet, this once.
            t.update(on=False, line="Not on the menu yet.")
        elif f["id"] == "doc":
            pass                                            # always
        elif f["id"] == "pics":
            if not c["spec"]["images"]:
                t.update(on=False, line="This one doesn't take pics.")
            elif not _pics_in(run):
                t.update(on=False, line="You'll need to send those separately.")
        elif f["id"] == "code":
            if not _emits(c, "html", "code"):
                t.update(on=False, line="You'll need to upload to your tool.")
        tiles.append(t)
    return tiles


@file_bp.route("/api/fillings", methods=["POST"])
@require_auth
def fillings():
    data = request.get_json() or {}
    c = CT.container(data.get("container", ""))
    if not c:
        return jsonify({"error": "No such container."}), 404
    return jsonify({"success": True, "tiles": menu(c, data.get("run", ""))})


# ---------------------------------------------------------------------------
# THE PARCEL — placeholders filled from the same facts as the terms.
# Shared with /api/parcel in app.py.
# ---------------------------------------------------------------------------

def parcel(c, data):
    facts = build_facts(c, data.get("form") or {})
    cp = data.get("copy") or {}
    out = {}
    for k, v in cp.items():
        if isinstance(v, str):
            out[k] = render_copy(c, v, facts)
        elif isinstance(v, list):
            out[k] = [({kk: render_copy(c, vv, facts) for kk, vv in x.items()} if isinstance(x, dict)
                       else render_copy(c, x, facts)) for x in v]
    name = next((facts[r["id"]] for g in c["needs"]["groups"] for r in g["rows"]
                 if r["type"] == "text" and facts.get(r["id"])), c["id"])
    slug = "".join(ch if ch.isalnum() else "-" for ch in str(name).lower()).strip("-") or c["id"]
    return {"facts": facts, "copy": out, "name": str(name), "slug": slug,
            "terms": render_terms(c, facts, data.get("chosen"))}


# ---------------------------------------------------------------------------
# THE DOC — the Hunch copy deck, drawn from the container.
# ---------------------------------------------------------------------------

GREY = RGBColor(0x80, 0x80, 0x80)
HUNCH_RED = "ED1C24"
PH_RED = RGBColor(0xFF, 0x00, 0x00)

def _shade(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)

def _width(cell, cm):
    cell.width = Cm(cm)

def _grid(table, left=4.0, right=11.9):
    """The deck's grid: 2263 / 6753 dxa. python-docx needs it said on the
    columns, the grid and every cell before Word or LibreOffice believes it."""
    table.autofit = False
    tbl = table._tbl
    grid = tbl.tblGrid
    for gc, cm in zip(grid.findall(qn("w:gridCol")), (left, right)):
        gc.set(qn("w:w"), str(int(Cm(cm).twips)))
    for col, cm in zip(table.columns, (left, right)):
        col.width = Cm(cm)
    for row in table.rows:
        cells = row.cells
        if cells[0]._tc is cells[-1]._tc:          # merged bar
            cells[0].width = Cm(left + right)
        else:
            cells[0].width = Cm(left); cells[1].width = Cm(right)

def _para(cell, text, bold=False, first=True, colour=None, align=None, before=6, after=6, line=1.2):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align: p.alignment = align
    # red brackets for anything unfilled — the deck's placeholder rule
    for i, chunk in enumerate(re.split(r"(\{\w+\}|\[[^\]]+\])", text or "")):
        if not chunk: continue
        r = p.add_run(chunk)
        r.bold = bold
        r.font.size = Pt(11); r.font.name = "Calibri"
        if colour is not None: r.font.color.rgb = colour
        if re.fullmatch(r"\{\w+\}|\[[^\]]+\]", chunk):
            r.font.color.rgb = PH_RED
            if chunk.startswith("{"): r.text = "[" + chunk[1:-1] + "]"
    return p

def _row(table, label, value, bold=False):
    cells = table.add_row().cells
    _width(cells[0], 4.0); _width(cells[1], 11.9)
    _para(cells[0], label)
    lines = [l for l in (value or "").split("\n")]
    for i, l in enumerate(lines):
        _para(cells[1], l, bold=bold, first=(i == 0), before=6 if i == 0 else 0,
              after=6 if i == len(lines) - 1 else 0)

def _bar(table, text, fill):
    cells = table.add_row().cells
    m = cells[0].merge(cells[1])
    _shade(m, fill)
    _para(m, text.upper(), bold=True, colour=RGBColor(0xFF, 0xFF, 0xFF),
          align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3, line=1.0)

_BOLD = ("subject", "headline", "button", "cta")

def _nice(module):
    s = module.replace("-", " ").replace("cta", "button")
    return s.strip().capitalize()

def _pretty(mod):
    """A fixed module's own words, from the spec's length column. Pointers
    to elsewhere (config.md, IMAGES) and furniture (logo, merge) aren't words."""
    L = mod.get("length", "").strip()
    if re.search(r"image|logo|merge|config\.md|see IMAGES", L, re.I):
        return ""
    return L

def _logo(c):
    """brandlook.md's Logo line names a file in the brand's assets/. PNG or JPG
    goes in; an SVG can't (docx won't take it) — the validator's job to say."""
    m = re.search(r"assets/([\w.\-]+\.(png|jpg|jpeg))", c["brand_data"].get("skin", {}).get("logo", ""), re.I)
    if not m:
        return None
    path = os.path.join(c["brand_data"]["folder"], "assets", m.group(1))
    return path if os.path.isfile(path) else None

def _mark(cell, path, words, height, align):
    p = cell.paragraphs[0]; p.alignment = align
    if path:
        p.add_run().add_picture(path, height=Cm(height))
    else:
        r = p.add_run(words); r.font.size = Pt(8); r.font.color.rgb = GREY; r.bold = True

def build_doc(c, P, tweaks=0, who=""):
    d = Document()
    client = c.get("client") or c["brand_data"].get("name", c["brand"])
    logo = _logo(c)
    for s in d.sections:
        s.left_margin = s.right_margin = Cm(2.54); s.top_margin = Cm(1.6); s.bottom_margin = Cm(2.0)
        # HAI2 left, the client's mark right — Hunch furniture holding the client's work.
        s.header.paragraphs[0].text = ""
        ht = s.header.add_table(rows=1, cols=2, width=Cm(15.9)); ht.autofit = False
        _mark(ht.cell(0, 0), HAI2, "ROBOT SANDWICH  \u00b7  HUNCH", 1.0, WD_ALIGN_PARAGRAPH.LEFT)
        _mark(ht.cell(0, 1), logo, client.upper(), 1.8, WD_ALIGN_PARAGRAPH.RIGHT)
        f = s.footer.paragraphs[0]; f.text = ""
        r = f.add_run(f"Supported by Hunch Robot Sandwich  \u00b7  {date.today().strftime('%-d %B %Y')}")
        r.font.size = Pt(8); r.font.color.rgb = GREY
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

    big = d.add_paragraph(); r = big.add_run("COPY DOC")
    r.font.name = "Bebas Neue"; r.font.size = Pt(40); r.font.color.rgb = RGBColor(0x1A, 0x19, 0x17)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Bebas Neue"); r._element.rPr.rFonts.set(qn("w:cs"), "Bebas Neue")
    big.paragraph_format.space_after = Pt(0); big.paragraph_format.line_spacing = 1.0
    t = d.add_paragraph(); r = t.add_run(f"{client.upper()} | {P['name']}")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = GREY

    meta = d.add_table(rows=0, cols=2); meta.style = "Table Grid"; meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in (("Client", client), ("Container", c.get("name", c["id"])),
                 ("Human in the loop", who or "\u2014"),
                 ("Date / version", date.today().strftime("%-d %B %Y") + " / v1"
                  + (f"  \u00b7  {tweaks} tweak{'s' if tweaks != 1 else ''} logged" if tweaks else ""))):
        cells = meta.add_row().cells; _width(cells[0], 4.0); _width(cells[1], 11.9)
        _para(cells[0], k, bold=True, before=0, after=0, line=1.15)
        _para(cells[1], v, before=0, after=0, line=1.15)
    d.add_paragraph()

    fill = HUNCH_RED                                   # the deck is Hunch furniture
    tb = d.add_table(rows=0, cols=2); tb.style = "Table Grid"; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    copy = P["copy"]
    mods = c["spec"]["modules"]
    groups = {m["module"]: m for m in mods if m.get("repeat")}
    for o in c["spec"]["outputs"]:
        _bar(tb, o.get("name") or o["id"], fill)
        for m in mods:
            if m.get("output") and m["output"] != o["id"]:
                continue
            mid, fb = m["module"], m.get("filled_by", "")
            if m.get("repeat"):
                for i, card in enumerate(copy.get(mid + "s") or [], 1):
                    for part, text in card.items():
                        _row(tb, f"{_nice(mid)} {i} — {_nice(part[len(mid)+1:] if part.startswith(mid + '-') else part)}",
                             text, bold=any(b in part for b in _BOLD))
                continue
            if any(mid.startswith(g + "-") for g in groups):
                continue
            if fb.startswith("writer"):
                v = copy.get(mid)
                if isinstance(v, list): v = "\n".join(v)
                _row(tb, _nice(mid), v or "", bold=any(b in mid for b in _BOLD))
            elif fb.startswith("assembled") or mid == "terms":
                _row(tb, _nice(mid), P["terms"])
            elif fb.startswith("fixed") and mid != "footer":   # the footer rides with the terms
                text = _pretty(m)
                if text:
                    _row(tb, _nice(mid), text, bold=any(b in mid for b in _BOLD))
    _grid(meta); _grid(tb)
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


# ---------------------------------------------------------------------------
# THE WRAP
# ---------------------------------------------------------------------------

def _zip(entries):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in entries:
            z.writestr(name, data)
    return buf.getvalue()


@file_bp.route("/api/wrap", methods=["POST"])
@require_auth
def wrap():
    data = request.get_json() or {}
    c = CT.container(data.get("container", ""))
    if not c:
        return jsonify({"error": "No such container."}), 404
    run = _clean(data.get("run"))
    if not run:
        return jsonify({"error": "No run id."}), 400
    try:
        P = parcel(c, data)
    except TermsError as e:
        return jsonify({"error": str(e)}), 400
    on = {t["id"] for t in menu(c, run) if t["on"]}
    take = [t for t in (data.get("take") or []) if t in on]
    if not take:
        return jsonify({"error": "Nothing's ticked."}), 400

    built = []                                           # (id, filename, bytes, line)
    if "doc" in take:
        built.append(("doc", f"{P['slug']}-copy.docx", build_doc(c, P, int(data.get("tweaks") or 0), session.get("name", "")),
                      "Every word, on the record"))
    if "pics" in take:
        folder = os.path.join(IMAGES_DIR, run)
        pics = [(n, open(os.path.join(folder, n), "rb").read()) for n in _pics_in(run)]
        built.append(("pics", f"{P['slug']}-pics.zip", _zip(pics),
                      f"{len(pics)} pic{'s' if len(pics) != 1 else ''}, as they came"))

    folder = os.path.join(WRAPS_DIR, run)
    os.makedirs(folder, exist_ok=True)
    for old in os.listdir(folder):
        os.remove(os.path.join(folder, old))
    files = []
    for fid, name, blob, line in built:
        with open(os.path.join(folder, name), "wb") as fh: fh.write(blob)
        files.append({"id": fid, "name": name, "line": line, "url": f"/api/wrap/{run}/{name}"})
    out = {"success": True, "files": files}
    if len(built) > 1:
        name = f"{P['slug']}-wrapped.zip"
        with open(os.path.join(folder, name), "wb") as fh:
            fh.write(_zip([(n, b) for _, n, b, _ in built]))
        out["all"] = {"name": name, "line": "The lot, in one", "url": f"/api/wrap/{run}/{name}"}
    return jsonify(out)


@file_bp.route("/api/wrap/<run>/<name>")
@require_auth
def take(run, name):
    run = _clean(run)
    if not run or not re.fullmatch(r"[a-z0-9-]+\.(docx|zip|pdf|html)", name or ""):
        return jsonify({"error": "No such file."}), 404
    return send_from_directory(os.path.join(WRAPS_DIR, run), name, as_attachment=True)
